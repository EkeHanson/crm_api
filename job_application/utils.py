import logging
import os
from pdfminer.high_level import extract_text
from docx import Document
from django.core.files.storage import default_storage
logger = logging.getLogger('job_applications')
import requests
import tempfile
import re

# Lazy initialization of SentenceTransformer
_model = None

import logging
from sentence_transformers import SentenceTransformer, util
import torch

logger = logging.getLogger('job_applications')

_model = None

def get_sentence_transformer_model():
    """Lazily load the SentenceTransformer model."""
    global _model
    if _model is None:
        try:
            _model = SentenceTransformer('all-MiniLM-L6-v2')
            device = torch.device('cpu')
            if _model.device.type == 'meta':
                _model.to_empty(device=device)
            else:
                _model.to(device)
            logger.info("Loaded SentenceTransformer model on CPU")
        except Exception as e:
            logger.exception(f"Failed to load SentenceTransformer model: {str(e)}")
            raise RuntimeError("Unable to initialize sentence transformer model")
    return _model

# def parse_resume(file_path):
#     """Extract text from PDF or DOCX files."""
#     try:
#         #logger.debug(f"Processing file_path: {file_path}")
#         temp_file_path = None
#         if file_path.startswith("http"):
#             response = requests.get(file_path)
#             if response.status_code != 200:
#                 logger.error(f"Failed to download file: {file_path}")
#                 return ""
#             temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
#             temp_file.write(response.content)
#             temp_file.close()
#             temp_file_path = temp_file.name
#             full_path = temp_file_path
#         else:
#             normalized_path = file_path.lstrip("/").replace("media/", "", 1)
#             logger.debug(f"Normalized path: {normalized_path}")
#             if default_storage.exists(normalized_path):
#                 full_path = default_storage.path(normalized_path)
#             else:
#                 logger.error(f"File does not exist: {normalized_path}")
#                 #print(f"File does not exist: {normalized_path}")
#                 return ""

#         #logger.debug(f"Full path: {full_path}")

#         ext = os.path.splitext(full_path)[1].lower()
#         if ext == '.pdf':
#             text = extract_text(full_path)
#         elif ext in ['.docx', '.doc']:
#             doc = Document(full_path)
#             text = '\n'.join([para.text for para in doc.paragraphs])
#         else:
#             logger.error(f"Unsupported file type: {ext}")
#             text = ""

#         # Clean up temporary file if created
#         if temp_file_path and os.path.exists(temp_file_path):
#             os.unlink(temp_file_path)
#         return text
#     except Exception as e:
#         logger.exception(f"Error parsing resume {file_path}: {str(e)}")
#         # Clean up temporary file on error
#         if temp_file_path and os.path.exists(temp_file_path):
#             os.unlink(temp_file_path)
#         return ""
    

def screen_resume(resume_text, job_description):
    """Compute similarity score between resume and job description."""
    try:
        if not resume_text or not job_description:
            logger.warning(f"Empty input: resume_text={bool(resume_text)}, job_description={bool(job_description)}")
            #print(f"Empty input: resume_text={bool(resume_text)}, job_description={bool(job_description)}")
            return 0.0
        model = get_sentence_transformer_model()
        from sentence_transformers import util
        resume_emb = model.encode(resume_text, convert_to_tensor=True)
        jd_emb = model.encode(job_description, convert_to_tensor=True)
        score = util.pytorch_cos_sim(resume_emb, jd_emb).item()
        return round(score * 100, 2)
    except Exception as e:
        logger.exception(f"Error screening resume: {str(e)}")
        return 0.0
    
import pdfplumber
import re
import os
import tempfile
import requests
from django.core.files.storage import default_storage
from docx import Document
import logging

logger = logging.getLogger('job_applications')

def parse_resume(file_path):
    try:
        logger.debug(f"Processing file_path: {file_path}")
        temp_file_path = None
        if file_path.startswith("http"):
            response = requests.get(file_path)
            if response.status_code != 200:
                logger.error(f"Failed to download file: {file_path}")
                return ""
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_file.write(response.content)
            temp_file.close()
            temp_file_path = temp_file.name
            full_path = temp_file_path
        else:
            normalized_path = file_path.lstrip("/").replace("media/", "", 1)
            logger.debug(f"Normalized path: {normalized_path}")
            if not default_storage.exists(normalized_path):
                logger.error(f"File does not exist: {normalized_path}")
                return ""
            full_path = default_storage.path(normalized_path)

        ext = os.path.splitext(full_path)[1].lower()
        text = ""
        if ext == '.pdf':
            try:
                with pdfplumber.open(full_path) as pdf:
                    text = '\n'.join([page.extract_text() or '' for page in pdf.pages])
            except Exception as e:
                logger.error(f"pdfplumber failed: {str(e)}. Falling back to pdfminer.")
                text = extract_text(full_path)
            # Clean OCR artifacts
            text = re.sub(r'^\d{1,2}/\d{1,2}/\d{2,4},\s*\d{1,2}/\d{2}\s*(?:AM|PM)\n?', '', text, flags=re.MULTILINE)
            text = re.sub(r'\s+', ' ', text).strip()  # Normalize whitespace
            logger.debug(f"Extracted resume text (first 1000 chars): {text[:1000]}")
        elif ext in ['.docx', '.doc']:
            doc = Document(full_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            logger.debug(f"Extracted resume text (first 1000 chars): {text[:1000]}")
        else:
            logger.error(f"Unsupported file type: {ext}")
            text = ""

        if temp_file_path and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        return text
    except Exception as e:
        logger.exception(f"Error parsing resume {file_path}: {str(e)}")
        if temp_file_path and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        return ""
    
# def extract_resume_fields(resume_text):
#     """Extract structured fields from resume text using regex."""
#     try:
#         extracted_data = {
#             "full_name": "",
#             "email": "",
#             "phone": "",
#             "qualification": "",
#             "experience": "",
#             "knowledge_skill": ""
#         }

#         # Extract full name (simplified: assume first line or capitalized words at start)
#         name_pattern = r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+'
#         name_match = re.search(name_pattern, resume_text, re.MULTILINE)
#         if name_match:
#             extracted_data["full_name"] = name_match.group(0).strip()

#         # Extract email
#         email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
#         email_match = re.search(email_pattern, resume_text)
#         if email_match:
#             extracted_data["email"] = email_match.group(0)

#         # Extract phone number
#         phone_pattern = r'\b(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
#         phone_match = re.search(phone_pattern, resume_text)
#         if phone_match:
#             extracted_data["phone"] = phone_match.group(0)

#         # Extract qualifications (e.g., degrees)
#         qual_pattern = r'\b(B\.Sc|Bachelor|M\.Sc|Master|Ph\.D|Diploma)\b.*?(?=\n|$|\b[A-Z])'
#         qual_matches = re.findall(qual_pattern, resume_text, re.IGNORECASE)
#         if qual_matches:
#             extracted_data["qualification"] = ", ".join(qual_matches).strip()

#         # Extract experience (e.g., years of experience or job titles with dates)
#         exp_pattern = r'(\d{1,2}\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|work))|((?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*(?:@|at)\s*(?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*\(\d{4}\s*-\s*(?:\d{4}|Present)\))'
#         exp_matches = re.findall(exp_pattern, resume_text, re.IGNORECASE)
#         if exp_matches:
#             experiences = []
#             for match in exp_matches:
#                 if match[0]:
#                     experiences.append(match[0])
#                 if match[1]:
#                     experiences.append(match[1])
#             extracted_data["experience"] = "; ".join(experiences).strip()

#         # Extract skills (e.g., list of common skills)
#         skills_pattern = r'\b(Python|Java|JavaScript|SQL|Project Management|Communication|Leadership|Teamwork|Problem Solving)\b'
#         skills_matches = re.findall(skills_pattern, resume_text, re.IGNORECASE)
#         if skills_matches:
#             extracted_data["knowledge_skill"] = ", ".join(set(skills_matches)).strip()

#         return extracted_data
#     except Exception as e:
#         logger.exception(f"Error extracting fields from resume: {str(e)}")
#         return {}
from datetime import datetime
import re
from dateutil.parser import parse as parse_date
from dateutil.relativedelta import relativedelta
import logging

logger = logging.getLogger('job_applications')

def extract_resume_fields(resume_text):
    try:
        extracted_data = {
            "full_name": "",
            "email": "",
            "phone": "",
            "qualification": "",
            "experience": [],
            "knowledge_skill": "",
            "employment_gaps": []
        }

        # Extract full name
        name_pattern = r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+'
        name_match = re.search(name_pattern, resume_text, re.MULTILINE)
        if name_match:
            extracted_data["full_name"] = name_match.group(0).strip()

        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, resume_text)
        if email_match:
            extracted_data["email"] = email_match.group(0)

        # Extract phone number
        phone_pattern = r'\b(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
        phone_match = re.search(phone_pattern, resume_text)
        if phone_match:
            extracted_data["phone"] = phone_match.group(0)

        # Extract qualifications
        qual_pattern = r'\b(B\.Sc|Bachelor|M\.Sc|Master|Ph\.D|Diploma)\b.*?(?=\n|$|\b[A-Z])'
        qual_matches = re.findall(qual_pattern, resume_text, re.IGNORECASE)
        if qual_matches:
            extracted_data["qualification"] = ", ".join(qual_matches).strip()

        # Extract experience with robust date pattern
        exp_pattern = r'((?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*(?:@|at)\s*(?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*\((?:(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\s*(\d{4}))\s*[-–—]\s*(?:(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\s*(\d{4}|Present))\))'
        exp_matches = re.findall(exp_pattern, resume_text, re.IGNORECASE | re.MULTILINE)
        job_entries = []
        for match in exp_matches:
            job_info, start_month, start_year, end_month, end_year = match
            start_date_str = f"{start_month or 'Jan'} {start_year}".strip()
            end_date_str = f"{end_month or 'Dec'} {end_year}".strip() if end_year != "Present" else "Present"
            job_title = job_info.split('(')[0].strip()
            job_entries.append({
                "job": job_title,
                "start_date": start_date_str,
                "end_date": end_date_str
            })
            logger.debug(f"Extracted job: {job_title}, Start: {start_date_str}, End: {end_date_str}")

        # Sort job entries by start date
        def parse_job_date(date_str):
            try:
                return parse_date(date_str, default=datetime(2000, 1, 1))
            except ValueError:
                return datetime.now() if date_str.lower() == "present" else datetime(2000, 1, 1)

        job_entries.sort(key=lambda x: parse_job_date(x["start_date"]))

        # Calculate employment gaps
        gap_threshold_months = 6
        gaps = []
        current_year = datetime.now().year

        for i in range(len(job_entries)):
            job = job_entries[i]
            try:
                start_date = parse_date(job["start_date"], default=datetime(int(job["start_date"].split()[-1]), 1, 1))
                end_date = parse_date(job["end_date"], default=datetime(int(job["end_date"].split()[-1]) if job["end_date"] != "Present" else current_year, 12, 31))
            except ValueError as e:
                logger.error(f"Error parsing dates for job {job['job']}: {str(e)}")
                continue

            extracted_data["experience"].append(f"{job['job']} ({job['start_date']} - {job['end_date']})")

            if i > 0:
                prev_job = job_entries[i - 1]
                try:
                    prev_end_date = parse_date(prev_job["end_date"], default=datetime(int(prev_job["end_date"].split()[-1]) if prev_job["end_date"] != "Present" else current_year, 12, 31))
                    gap_months = relativedelta(start_date, prev_end_date).months + relativedelta(start_date, prev_end_date).years * 12
                    if gap_months > gap_threshold_months:
                        gaps.append({
                            "gap_start": prev_end_date.strftime("%Y-%m"),
                            "gap_end": start_date.strftime("%Y-%m"),
                            "duration_months": gap_months
                        })
                        logger.debug(f"Detected gap: {prev_end_date.strftime('%Y-%m')} to {start_date.strftime('%Y-%m')} ({gap_months} months)")
                except ValueError as e:
                    logger.error(f"Error calculating gap for job {job['job']}: {str(e)}")
                    continue

        extracted_data["employment_gaps"] = gaps

        # Extract skills
        skills_pattern = r'\b(Python|Java|JavaScript|SQL|Project Management|Communication|Leadership|Teamwork|Problem Solving|HuggingFace|Transformers|SpaCy|NLTK|PyTorch|TensorFlow|AWS|GCP)\b'
        skills_matches = re.findall(skills_pattern, resume_text, re.IGNORECASE)
        if skills_matches:
            extracted_data["knowledge_skill"] = ", ".join(set(skills_matches)).strip()

        logger.debug(f"Extracted resume data: {extracted_data}")
        return extracted_data
    except Exception as e:
        logger.exception(f"Error extracting fields from resume: {str(e)}")
        return {}